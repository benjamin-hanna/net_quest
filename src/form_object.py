import xlsxwriter
import os
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

class FormObject():

    jaa_count = 0

    def __init__(self, output_directory, network_name, main, rest, loop_count):
        self.output_directory = output_directory
        self.network_name = network_name
        self.main = main
        self.rest = rest
        self.loop_count = loop_count

    def network_form_switcher(self):
        network_name = self.network_name
        default = "Intake form not available."
        getattr(self, 'case_' + str(network_name), lambda: default)()
    
    def case_anthem(self):
        if FormObject.jaa_count < 1:
            self.case_jaa()
            FormObject.jaa_count += 1

        else:
            pass
    
    
    def case_premera(self):
        if FormObject.jaa_count < 1:
            self.case_jaa()
            FormObject.jaa_count += 1

        else: 
            pass
    
    def case_empire(self):
        if FormObject.jaa_count < 1:
            self.case_jaa()
            FormObject.jaa_count += 1

        else: 
            pass
    
    def case_jaa(self):
        main = self.main

        file_name = "_".join(['CH.JAAForm' , main[0][0][1] ,'[DATE]' , '.xlsx' ])
        output_directory = Path(os.environ['HOME'] + '/Desktop/output/intake_forms/' + file_name)

        workbook = xlsxwriter.Workbook(output_directory)
        worksheet = workbook.add_worksheet()

        # Creates four cell formats
        cell_format1 = workbook.add_format({
            'bold': '1',
            'font_size': '10',
            'font_name': 'Arial'})

        cell_format2 = workbook.add_format({
            'font_size': '10',
            'font_name': 'Arial'})

        cell_format3 = workbook.add_format({
            'bold': '1',
            'font_size': '10',
            'font_name': 'Arial',
            'font_color': 'Blue'})

        merge_format = workbook.add_format({
            'bold': 1,
            'border': 1,
            'align': 'center',
            'valign': 'vcenter',
            'fg_color': '#119922',
            'font_name': 'Arial'})

        # Merges and formats header
        worksheet.set_column('A:B', 44)
        worksheet.merge_range('A1:B1', 'Collective Health', merge_format)

        # Write first column (fields) for JAA intake form
        worksheet.write(4, 0, 'Employer Name', cell_format1)
        worksheet.write(5, 0, 'Employer Location / Headquarter', cell_format1)
        worksheet.write(6, 0, 'Business (SIC if available)', cell_format1)
        worksheet.write(7, 0, 'Number of Employees', cell_format1)
        worksheet.write(8, 0, 'Number of Covered Employees', cell_format1)
        worksheet.write(9, 0, 'Number of total Members (employees and dependents)', cell_format1)
        worksheet.write(10, 0, 'Census', cell_format1)
        worksheet.write(11, 0, 'Current PPO/HMO Network', cell_format1)
        worksheet.write(12, 0, 'Current Administrator (Carrier & TPA)', cell_format1)
        worksheet.write(13, 0, 'Proposed Network (CA Only or JAA *)', cell_format1)
        worksheet.write(14, 0, 'Number of out-of-state Employees (If requesting JAA)', cell_format1)
        worksheet.write(15, 0, 'Proposed Effective Date', cell_format1)
        worksheet.write(16, 0, 'Broker or Consultant Name', cell_format1)
        worksheet.write(17, 0, 'Broker or Consultant Location', cell_format1)

        # Writes JAA callout
        worksheet.write(20, 0, '* JAA includes BlueCard access', cell_format3)

        # Write second column (values) for JAA intake form using global variables from render_text.py
        worksheet.write(4, 1, str(main[0][0][1]), cell_format2) # prospect_name
        hq = " ".join([ main[0][2][1], main[0][5][1], main[0][3][1], main[0][4][1]])
        worksheet.write(5, 1, hq, cell_format2) # hq, city, state, zip code
        worksheet.write(6, 1, ' ', cell_format2)
        worksheet.write(7, 1, str(main[0][13][1]), cell_format2) # ees
        worksheet.write(8, 1, str(main[0][14][1]), cell_format2) # covered es

        worksheet.write(9, 1, 'See Census', cell_format2) # total members
        worksheet.write(10, 1, 'Attached', cell_format2) # see census

        if main[0][17][1] == True: # Kaiser
            worksheet.write(11, 1, 'Kaiser', cell_format2)
        else: worksheet.write(11, 1, '', cell_format2)

        worksheet.write(12, 1, str(main[0][15][1]), cell_format2) # current medical carrier
        worksheet.write(13, 1, 'JAA', cell_format2) # proposed network
        worksheet.write(14, 1, 'See Census', cell_format2) # out of state employees
        worksheet.write(15, 1, str(main[0][6][1]), cell_format2) # start date
        worksheet.write(16, 1, str(main[0][10][1]), cell_format2) # broker contact
        worksheet.write(17, 1, str(main[0][9][1]), cell_format2) # broker address

        workbook.close()

    def case_blue_shield_ca(self):
        main = self.main
        rest = self.rest
        loop_count = self.loop_count

        bsc_form = "src/assets/intake_forms/bsc/Shared Advantage Quote Intake Form- 01.2020.docx"
        file_name = "_".join(['CH.Shared Advantage Quote Intake Form' , main[0][0][1] , '.docx' ])
        output_directory = Path(os.environ['HOME'] + '/Desktop/output/intake_forms/' + file_name)

        document = Document(bsc_form)

        now = datetime.now()
        year = now.strftime("%Y")
        month = now.strftime("%m")
        day = now.strftime("%d")
        array = []
        array.append(day)
        array.append(month)
        array.append(year)
        date = ("/".join(array))

        requestor_info = document.tables[1]
        requestor_info.cell(0, 0).text = "  TPA: Collective Health"
        requestor_info.cell(0, 1).text = "  TPA Salesperson: ", main[0][1][1] # rvp
        requestor_info.cell(0, 2).text = "  Date Submitted: ", date
        requestor_info.cell(1, 0).text = "  Due Date to TPA: ", main[0][7][1] # due date
        requestor_info.cell(1, 2).text = "  Group Effective Date: ", main[0][6][1] # start date

        if rest[loop_count][7][1]:  # zip code discount analysis
            FormObject.check_a_box(document, True, 1, (2, 0))
        else:
            FormObject.check_a_box(document, False, 1, (2, 1))

        if rest[loop_count][4][1]: # geo access analysis
            FormObject.check_a_box(document, True, 1, (1, 5))
        else:
            FormObject.check_a_box(document, False, 1, (1, 5))

        if rest[loop_count][9][1]: # questionnaire
            FormObject.check_a_box(document, True, 1, (2, 3))
        else:
            FormObject.check_a_box(document, False, 1, (2, 4))

        requestor_info.cell(4, 0).text = "  Blue Shield of CA Salesperson Name: TBD" # BSC sales person

        prospect_broker = document.tables[2]

        prospect_broker.cell(0, 0).text ="  Prospect Name: ", main[0][0][1] # prospect
        
        prospect_address = " ".join(["  Group's Corporate Headquarters (Complete Address): ", main[0][2][1], main[0][5][1], main[0][3][1], main[0][4][1]]) # hq, city, state, zip code
        prospect_broker.cell(1, 0).text = prospect_address 

        prospect_broker.cell(2, 0).text ="  Broker Firm Name: ", main[0][8][1] # broker
        prospect_broker.cell(2, 2).text =" Individual Broker Name: ", main[0][10][1] # broker contact
        prospect_broker.cell(3, 0).text = "  Broker Firm Complete Address: ", main[0][9][1] # broker address
        prospect_broker.cell(4, 0).text = "  Broker Contact Number: ", main[0][11][1] # broker phone 
        prospect_broker.cell(4, 1).text = "  Eligible Employees:  ", main[0][13][1] # ees
        prospect_broker.cell(4, 3).text = "  Covered Employees:  ", main[0][14][1] # ces

        # if main[0][17][1]: # kaiser
        #     FormObject.check_a_box(document, True, 2, (4, 8))
        # else:
        #     FormObject.check_a_box(document, False, 2, (4, 8))

        FormObject.check_a_box(document, True, 2, (6, 0)) # Quote Out of State Employees

        FormObject.check_a_box(document, False, 2, (7, 0)) # Is this a tribal account requiring Medicare Like Pricing (MLR):

        if main[0][16][1] == "self-funded": # current medical funding
            FormObject.check_a_box(document, True, 2, (8, 0))
        else:
            FormObject.check_a_box(document, False, 2, (8, 0))

        prospect_broker.cell(10, 0).text ="Current Carrier: ", main[0][15][1] # current med carrier 

        document.save(output_directory)

    @classmethod
    def make_box_checked(cls):
        element = OxmlElement('w:checked')
        element.set(qn('w:val'),"true")
        return element

    @classmethod
    def check_a_box(cls, document, is_checked, table_index, coordinates):

        check_boxes= document.tables[table_index].cell(coordinates[0], coordinates[1])._element.xpath('.//w:checkBox')
        if is_checked:
            index = 0

        else:
            index = 1

        check_boxes[index].append(FormObject.make_box_checked())

    def case_delta(self):
        main = self.main

        delta_form = "src/assets/intake_forms/delta/[DATE]_DeltaRequestForm_[Company].docx"
        document = Document(delta_form)

        header_name = " ".join([main[0][0][1] , main[0][7][1]]) # prospect, due date

        document.paragraphs[0].text = header_name # prospect header

        table_0 = document.tables[0]
        table_0.cell(0, 1).text = main[0][0][1] # prospect 
        table_0.cell(1, 1).text = main[0][2][1]  # prospect hq
        table_0.cell(2, 1).text = " ".join([main[0][5][1], main[0][3][1], main[0][4][1]]) # city, state, zip
        table_0.cell(3, 1).text = main[0][13][1]  # ees

        table_1 = document.tables[1]
        table_1.cell(0, 1).text = main[0][10][1]  # broker contact
        table_1.cell(1, 1).text = main[0][9][1] # broker address

        table_2 = document.tables[2]
        table_2.cell(0, 1).text = main[0][6][1] # start date

        file_name = "_".join(['CH.Delta Intake Form' , main[0][0][1] , '.docx' ])
        output_directory = Path(os.environ['HOME'] + '/Desktop/output/intake_forms/' + file_name)

        document.save(output_directory)

    def case_esi(self):
        main = self.main
        rest = self.rest
        loop_count = self.loop_count

        esi_form = "src/assets/intake_forms/esi/CH.ESI_Request Form.docx"
        file_name = "_".join(['CH.ESI_Request Form' , main[0][0][1] , '.docx' ])
        output_directory = Path(os.environ['HOME'] + '/Desktop/output/intake_forms/' + file_name)

        document = Document(esi_form)

        esi_table = document.tables[0]

        esi_table.cell(0, 1).text = main[0][0][1] # prospect
        esi_table.cell(1, 1).text = main[0][8][1] # broker
        esi_table.cell(2, 1).text = main[0][6][1] # start date
        esi_table.cell(3, 1).text = " ".join([main[0][15][1], main[0][18][1]]) # incumbent med carrier, pbm
        esi_table.cell(4, 1).text = main[0][13][1] # ees
        esi_table.cell(5, 1).text = "See attached plan designs"

        if rest[loop_count][6][1]: # claims repricing
            esi_table.cell(9, 1).text = 'Yes'
        else:
            esi_table.cell(9, 1).text = 'No'

        if rest[loop_count][5][1]: # disruption analysis
            esi_table.cell(13, 1).text = 'Yes'
        else:
            esi_table.cell(13, 1).text = 'No'

        if rest[loop_count][9][1]: # questionnaire 
            esi_table.cell(15, 1).text = 'Yes'
        else:
            esi_table.cell(15, 1).text = 'No'

        document.save(output_directory)

