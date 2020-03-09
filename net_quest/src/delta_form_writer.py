import os
from pathlib import Path
from docx import Document


def edit_delta_form(global_items):

    delta_form = Path(os.environ['HOME'] + '/bin/net_quest/src/assets/intake_forms/delta/[DATE]_DeltaRequestForm_[Company].docx')

    document = Document(delta_form)

    prospect = global_items[0][1]

    due_date = global_items[7][1]

    header_name = " ".join([prospect , due_date])

    document.paragraphs[0].text = header_name # prospect header

    table_0 = document.tables[0]
    table_0.cell(0, 1).text = global_items[0][1] # prospect name
    table_0.cell(1, 1).text = global_items[2][1]  # prospect hq
    table_0.cell(3, 1).text = global_items[13][1]  # ees

    table_1 = document.tables[1]
    table_1.cell(0, 1).text = global_items[10][1]  # broker contact
    table_1.cell(1, 1).text =  global_items[9][1] # broker address

    table_2 = document.tables[2]
    table_2.cell(0, 1).text = global_items[6][1] # start date

    file_name = "_".join(['CH.Delta Intake Form' , prospect , '.docx' ])
    output_directory = Path(os.environ['HOME'] + '/Desktop/net_quest_output/intake_forms/' + file_name)


    document.save(output_directory)
