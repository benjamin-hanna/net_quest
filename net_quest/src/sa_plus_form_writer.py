import os
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


def edit_sa_plus_form(global_items, filtered_dictionary):

    for dictionary_list in filtered_dictionary['network_info']:

        # global variables
        prospect = global_items[0][1]
        rvp = global_items[1][1]
        current_med_carrier = global_items[15][1]

        # stoploss
        policy_period = global_items[22][1],
        specific_deductible = global_items[23][1],
        stoploss_aggregate_corridor = global_items[24][1],

        # instance variables
        is_geo_access = dictionary_list[0]['is_geo_access']
        is_zip_discount_analysis = dictionary_list[0]['is_zip_discount_analysis']
        is_questionnaire = dictionary_list[0]['is_questionnaire']
        is_stop_loss = dictionary_list[0]['is_stop_loss']

    bsc_form = Path(os.environ['HOME'] + '/bin/net_quest/src/assets/intake_forms/bsc/Shared Advantage Quote Intake Form- 2019 .docx')
    file_name = "_".join(['CH.Shared Advantage Quote Intake Form' , prospect , '.docx' ])
    
    output_directory = Path(os.environ['HOME'] + '/Desktop/net_quest_output/intake_forms/' + file_name)

    document = Document(bsc_form)

    now = datetime.now()
    year = now.strftime("%Y")
    month = now.strftime("%m")
    day = now.strftime("%d")

    array = []
    array.append(day)
    array.append(month)
    array.append(year)

    date = ("/".join(array))

    requestor_info = document.tables[1]

    requestor_info.cell(0, 0).text = "  TPA: Collective Health"

    requestor_info.cell(0, 1).text = "  TPA Salesperson: ", rvp

    requestor_info.cell(0, 2).text = "  Date Submitted: ", date

    requestor_info.cell(1, 0).text = "  Due Date to TPA: ", global_items[20][1]

    requestor_info.cell(1, 2).text = "  Group Effective Date: ", global_items[7][1]


    if is_zip_discount_analysis:  # zip code discount analysis
        check_a_box(document, True, 1, (2, 0))
    else:
        check_a_box(document, False, 1, (2, 1))

    if is_geo_access: # geo access analysis
        check_a_box(document, True, 1, (1, 5))
    else:
        check_a_box(document, False, 1, (1, 5))

    if is_questionnaire: # questionnaire
        check_a_box(document, True, 1, (2, 3))
    else:
        check_a_box(document, False, 1, (2, 4))


    requestor_info.cell(4, 0).text = "  Blue Shield of CA Salesperson Name: TBD" # BSC sales person

    check_a_box(document, True, 1, (5, 0)) # Shared Advantage Plus Quote


    if is_stop_loss:
        check_a_box(document, True, 1, (5, 2))
    else: pass


    check_a_box(document, True, 1, (5, 3)) # Build Blue Card access fees into PEPM SA Fee

    check_a_box(document, False, 1, (7, 0)) # Tiered Network

    check_a_box(document, True, 1, (8, 0)) # Shield Savings


    prospect_broker = document.tables[2]

    prospect_broker.cell(0, 0).text ="  Prospect Name: ", global_items[0][1] # prospect name

    prospect_broker.cell(1, 0).text ="  Group's Corporate Headquarters (Complete Address): ", global_items[2][1] # prospect hq

    prospect_broker.cell(2, 0).text ="  Broker Firm Name: ", global_items[8][1] # brokerage

    prospect_broker.cell(2, 2).text =" Individual Broker Name: ", global_items[10][1] # broker contact

    prospect_broker.cell(3, 0).text = "  Broker Firm Complete Address: ", global_items[9][1] # broker address

    prospect_broker.cell(4, 0).text = "  Broker Contact Number: ", global_items[11][1] # broker phone number

    prospect_broker.cell(4, 1).text = "  Eligible Employees:  ", global_items[13][1] # ees

    prospect_broker.cell(4, 3).text = "  Covered Employees:  ", global_items[14][1] # ces

    if global_items[17][1]: # kaiser
        check_a_box(document, True, 2, (4, 4))
    else:
        check_a_box(document, False, 2, (4, 4))

    check_a_box(document, True, 2, (6, 0)) # Quote Out of State Employees

    check_a_box(document, False, 2, (7, 0)) # Is this a tribal account requiring Medicare Like Pricing (MLR):

    if global_items[16][1] == "Self-funded": # current medical funding
        check_a_box(document, True, 2, (8, 0))
    else:
        check_a_box(document, False, 2, (8, 0))

    prospect_broker.cell(10, 0).text ="Current Carrier: ", current_med_carrier 

    current_stoploss_info = document.tables[3]

    # if global_items[17][1] != "":
    #     check_a_box(document, True, 3, (0, 0))

    requested_stoploss_info = document.tables[4]


    document.save(output_directory)


def make_box_checked():
    element = OxmlElement('w:checked')
    element.set(qn('w:val'),"true")
    return element

def check_a_box(document, is_checked, table_index, coordinates):

    check_boxes= document.tables[table_index].cell(coordinates[0], coordinates[1])._element.xpath('.//w:checkBox')
    if is_checked:
        index = 0

    else:
        index = 1

    check_boxes[index].append(make_box_checked())
